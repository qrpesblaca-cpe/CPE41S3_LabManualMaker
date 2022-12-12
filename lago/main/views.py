import base64, os, re, requests, shutil
from .forms import RegistrationForm, LabManualForm, CourseCodeForm, CourseTitleForm, UpdateProfileForm
from .models import labmanual, course_code_db, course_title_db
from django.contrib import messages
from django.contrib.auth import login,logout, authenticate, update_session_auth_hash
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import AuthenticationForm, PasswordChangeForm
from django.http import HttpResponse
from django.shortcuts import  render, redirect
from django.utils.html import strip_tags
from django.utils.safestring import SafeString
from django.views.generic import ListView, DetailView
from docx import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from htmldocx import HtmlToDocx
from PIL import Image
from io import BytesIO

# View existing lab manuals
# -----------------------------
class labListView(ListView):
    model = labmanual
    labmanual_list = labmanual.objects.all()
    template_name = 'main/view.html'
    context_object_name: 'labmanual_list'
    ordering = ['-date_created']

# Preview of lab manuals
# -----------------------------
class labDetailView(DetailView):
    model = labmanual
    labmanual_list = labmanual.objects.all()
    context_object_name: 'labmanual_list'
    template_name = 'main/preview.html'

# Update lab manuals
# -----------------------------
@login_required(login_url='/')
def update(request, pk):
    lab = labmanual.objects.get(id=pk)
    form = LabManualForm(request.POST or None,instance=lab)
    if form.is_valid():
        form.save()
        messages.success(request,"Laboratory manual updated successfully!", extra_tags='invalid')
        return redirect('/home/view/')
    return render(request, 'main/update.html',{'form':form})

# Requires user to sign up 
# before accessing /home
# -----------------------------
@login_required(login_url='/')
def home(response):
    return render(response, "main/home.html", {})

# Inserts data 
# from forms into the labmanual 
# model
# -----------------------------
@login_required(login_url='/')
def insertlab(request):
    if request.method == "POST":
        labmanual_form = LabManualForm(request.POST, request.FILES)
        if labmanual_form.is_valid():
            labmanual_form.instance.author = request.user
            labmanual_form.save()
            messages.success(request, 'Laboratory manual successfully created!')
            return redirect('/home/view/')
        else:
            messages.error(request, "Error: Laboratory manual creation failed.", extra_tags='invalid')
            return redirect('/home/')
    labmanual_form = LabManualForm()
    course_coded = course_code_db.objects.all()
    return render(request=request, template_name="main/insertlab.html", context={'labmanual_form':labmanual_form,'course_code_db':course_coded})
    
# About the team
# --------------------------
@login_required(login_url='/')
def about(response):
    return render(response, "main/about.html", {})

# user Profile
# --------------------------
@login_required(login_url='/')
def userProfile(response):
    return render(response, "main/profile.html", {})

# Add new course
# --------------------------
@login_required(login_url='/')
def addCourse(request):
    if request.method == "POST":
        course_code_form = CourseCodeForm(request.POST)
        course_title_form = CourseTitleForm(request.POST)
        if course_code_form.is_valid():
            course_code_form.save()
            if course_title_form.is_valid():
                course_title_form.instance.code = course_code_db.objects.latest('id')
                course_title_form.save()
                messages.success(request, 'Course successfully added!')
        return redirect('/home/view/')
    course_code_form = CourseCodeForm()
    course_title_form = CourseTitleForm()
    return render(request=request, template_name="main/addcourse.html", context={'course_code_form':course_code_form, 'course_title_form':course_title_form})

# User creation
# -----------------------------
def signup(request):
	if request.method == "POST":
		form = RegistrationForm(request.POST)
		if form.is_valid():
			user = form.save()
			login(request, user)
			messages.success(request, "Signup successful.", extra_tags='valid')
			return redirect('/')
		messages.error(request, "Error: User Registration failed.", extra_tags='invalid')
	form = RegistrationForm()
	return render(request, 'main/signup.html',{'register_form':form})

# Change Password
# --------------------------
@login_required(login_url='/')
def change_password(request):
    if request.method == 'POST':
        form = PasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()
            update_session_auth_hash(request, user)
            messages.success(request, 'Your password was successfully updated!')
            return redirect("home")
        else:
            messages.error(request, 'Please correct the error below.')
    else:
        form = PasswordChangeForm(request.user)
    return render(request, 'main/changepass.html', {
        'form': form
    })

# Update information
# --------------------------
@login_required(login_url='/')
def update_profile(request):
    form = UpdateProfileForm(request.POST or None, instance=request.user)
    if form.is_valid():
        form.save()
        messages.success(request,"Profile updated successfully!", extra_tags='invalid')
        return redirect('/home/view/')
    else:
        messages.error(request,"Error: Profile update failure!", extra_tags='invalid')
        form = UpdateProfileForm(instance=request.user)
        args = {'form':form}
        return render(request, 'main/updateprofile.html',args)

# User sign in
# ---------------------------
def signin(request):
	if request.method == "POST":
		form = AuthenticationForm(request, data=request.POST)
		if form.is_valid():
			username = form.cleaned_data.get('username')
			password = form.cleaned_data.get('password')
			user = authenticate(username=username, password=password)
			if user is not None:
				login(request, user)
				#messages.info(request, f"You are now logged in as {username}.")
				return redirect("home")
			else:
				messages.error(request,"Error: Invalid username or password.", extra_tags='invalid')
		else:
			messages.error(request,"Error: Invalid username or password.", extra_tags='invalid')
	form = AuthenticationForm()
	return render(request=request, template_name="main/signin.html", context={"login_form":form})

# User sign out
# ---------------------------
def signout(request):
	logout(request)
	return redirect("/")

# Download lab manual template
# ---------------------------
@login_required(login_url='/')
def downloadTemp(request, id):
    act_no, lab_title, course_code, course_title, objectives, ilos, discussion, res, procedures, questions, supplementary, image_1, image_2 = getLab(id)


    # Create document
    # -----------------------------
    document = Document()
    docx_title = course_code + " - Activity No." + str(act_no) + ".docx"
    core_properties = document.core_properties

    # Insert author & comment
    # -----------------------------
    core_properties.author = 'Laboratory Manual On the Go'
    core_properties.comments = 'created by Laboratory Manual on the Go'

    # Create table
    # -----------------------------
    table = document.add_table(rows=27, cols=2)
    table.style = 'Table Grid'
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(12)
    font.bold = False

    downloadImage(procedures)

    image_files = os.listdir('temp_images')

    # Merge specific rows
    # -----------------------------
    table.cell(0,0).merge(table.cell(0,1))
    table.cell(1,0).merge(table.cell(1,1))
    x = 6
    for x in range(6,26):
        table.cell(x,0).merge(table.cell(x,1))

    # Insert labels and content into cells
    # ------------------------------------
    table.rows[0].cells[0].paragraphs[0].add_run('Activity No. ' + str(act_no)).bold = True
    table.rows[0].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.rows[1].cells[0].paragraphs[0].add_run(lab_title).bold = True
    table.rows[1].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.rows[2].cells[0].paragraphs[0].add_run('Course Code: ' + course_code).bold = True
    table.rows[3].cells[0].paragraphs[0].add_run('Course Title: ' + course_title).bold = True
    table.rows[4].cells[0].paragraphs[0].add_run('Section: ').bold = True
    table.rows[5].cells[0].paragraphs[0].add_run('Name/s: ').bold = True

    table.rows[2].cells[1].paragraphs[0].add_run('Program: ').bold = True
    table.rows[3].cells[1].paragraphs[0].add_run('Date Performed: ').bold = True
    table.rows[4].cells[1].paragraphs[0].add_run('Date Submitted: ').bold = True
    table.rows[5].cells[1].paragraphs[0].add_run('Instructor: Engr. ' + request.user.first_name + ' ' + request.user.last_name + '\n\n').bold = True

    table.rows[6].cells[0].paragraphs[0].add_run('1. Objective: ').bold = True
    table.rows[7].cells[0].paragraphs[0].add_run(objectives).bold = False
    table.rows[8].cells[0].paragraphs[0].add_run('2. Intended Learning Outcomes (ILOs): ').bold = True
    table.rows[9].cells[0].paragraphs[0].add_run(ilos).bold = False
    table.rows[10].cells[0].paragraphs[0].add_run('3. Discussion: ').bold = True
    table.rows[11].cells[0].paragraphs[0].add_run(discussion).bold = False
    table.rows[11].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    table.rows[12].cells[0].paragraphs[0].add_run('4. Resources: ').bold = True
    table.rows[13].cells[0].paragraphs[0].add_run(res).bold = False
    table.rows[14].cells[0].paragraphs[0].add_run('5. Procedures: ').bold = True
    table.rows[16].cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.rows[17].cells[0].paragraphs[0].add_run('6. Results: ').bold = True
    # table.rows[17].cells[0].paragraphs[0].add_run(image_2.url)
    table.rows[19].cells[0].paragraphs[0].add_run('7. Observations: ').bold = True
    # table.rows[19].cells[0].paragraphs[0].add_run(image_1.url)
    table.rows[21].cells[0].paragraphs[0].add_run('8. Questions: ').bold = True
    table.rows[22].cells[0].paragraphs[0].add_run(questions).bold = False
    table.rows[23].cells[0].paragraphs[0].add_run('9. Conclusions: ').bold = True
    table.rows[25].cells[0].paragraphs[0].add_run('10. Supplementary Activity: ').bold = True
    table.rows[26].cells[0].paragraphs[0].add_run(supplementary).bold = False
    
    for image_file in image_files:
        table.rows[16].cells[0].paragraphs[0].add_run().add_picture('temp_images/' + image_file)

    # Parse HTML to Docx
    parser = HtmlToDocx()
    procedures = re.sub(r'<img.*?>', '', procedures)
    html = procedures
    parser.add_html_to_cell(html, table.cell(0,30))

    # Prepare document for download        
    # -----------------------------
    f = BytesIO()

    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response

# Fetch lab manual data from the database        
# -----------------------------
def getLab(id):
    lab_specific = labmanual.objects.all().filter(id=id)
    for lab in lab_specific:
        act_no = lab.act_no
        lab_title = lab.lab_title
        course_code = lab.course_code
        course_title =lab.course_title
        objectives = lab.objective
        ilos = lab.ilos
        discussion = lab.discussion
        res = lab.res
        procedures = lab.procedures
        questions = lab.questions
        supplementary = lab.supplementary
        image_1 = lab.image_1
        image_2 = lab.image_2
        return act_no, lab_title, course_code, course_title, objectives, ilos, discussion, res, procedures, questions, supplementary, image_1, image_2

# Delete data from the database
# -----------------------------
def deleteLab(request,id):
    labmanual.objects.filter(id=id).delete()
    messages.success(request,"Successfully deleted!", extra_tags='invalid')
    return redirect("/home/view/")


def downloadImage(html):
    shutil.rmtree('temp_images', ignore_errors=True)
    img_tags = re.findall(r'<img.*?>', html)
    img_urls = []
    for img_tag in img_tags:
        src_attr = re.search(r'src="(.*?)"', img_tag)
        img_tag = img_tag[:-1] + ' width="100" height="100">'
        if src_attr:
            img_urls.append(src_attr.group(1))

    temp_dir = 'temp_images'
    if not os.path.exists(temp_dir):
        os.mkdir(temp_dir)

    for url in img_urls:
        response = requests.get(url)
        with open(os.path.join(temp_dir, os.path.basename(url)), 'wb') as f:
            f.write(response.content)

    for filename in os.listdir('temp_images'):
        with Image.open(os.path.join('temp_images', filename)) as img:
            img = img.resize((200, 200))
            img.save(os.path.join('temp_images', filename))
    
