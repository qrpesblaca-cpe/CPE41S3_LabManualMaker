from django.http import HttpResponse
from docx import *
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date
from io import BytesIO
import docx


def downloadTemp(request):

    document = Document()
    docx_title = "LAGO-test.docx"

    #Create Table
    # -----------------------------
    table = document.add_table(rows=26, cols=2)
    table.style = 'Table Grid'
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(12)
    font.bold = True

    #Merge Specific Rows
    # -----------------------------
    table.cell(0,0).merge(table.cell(0,1))
    table.cell(1,0).merge(table.cell(1,1))
    x = 6
    for x in range(6,26):
        table.cell(x,0).merge(table.cell(x,1))

    p=table.columns[0].cells[0].add_paragraph('Activity No. ')
    p.paragraph_format.space_after = Pt(12)
    p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER

    table.columns[0].cells[2].text = 'Course Code: '
    table.columns[0].cells[3].text = 'Course Title: '
    table.columns[0].cells[4].text = 'Section: '
    table.columns[0].cells[5].text = 'Name/s: '

    table.columns[1].cells[2].text = 'Program: '
    table.columns[1].cells[3].text = 'Date Performed: '
    table.columns[1].cells[4].text = 'Date Submitted: '
    table.columns[1].cells[5].text = 'Instructor: \n\n'

    table.columns[0].cells[6].text = '1. Objective:'
    table.columns[0].cells[8].text = '2. Intended Learning Outcomes (ILOs):'
    table.columns[0].cells[10].text = '3. Discussion:'
    table.columns[0].cells[12].text = '4. Resources:'
    table.columns[0].cells[14].text = '5. Procedures:'
    table.columns[0].cells[16].text = '6. Results:'
    table.columns[0].cells[18].text = '7. Observations:'
    table.columns[0].cells[20].text = '8. Questions:'
    table.columns[0].cells[22].text = '9. Conclusions:'
    table.columns[0].cells[24].text = '10. Supplementary Activity:'

    # Prepare document for download        
    # -----------------------------
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response